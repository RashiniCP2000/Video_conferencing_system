import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def create_diary_docx():
    md_path = r"c:\Users\Rashini\Desktop\video conferencing system\MeetNova_Internship_Daily_Diary.md"
    docx_path = r"c:\Users\Rashini\Desktop\video conferencing system\MeetNova_Internship_Daily_Diary.docx"

    if not os.path.exists(md_path):
        print(f"Error: Markdown file not found at {md_path}")
        return

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = docx.Document()

    # Define clean, professional document styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x1e, 0x29, 0x3b) # Slate 800

    # Add Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Internship Daily Task Diary")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8) # Royal Blue

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("MeetNova Video Conferencing & Collaboration System")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69) # Slate 600

    # Horizontal Divider Line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_div = p_div.add_run("____________________________________________________")
    run_div.font.color.rgb = RGBColor(0xe2, 0xe8, 0xf0)

    # Document details block
    doc.add_paragraph() # Spacer
    details = [
        ("Duration:", " March 23, 2026 – July 4, 2026 (15 Weeks)"),
        ("Role:", " Software Engineering Intern"),
        ("Project:", " MeetNova — A secure MERN-based video conferencing and real-time collaboration suite.")
    ]
    for label, val in details:
        p = doc.add_paragraph()
        run_l = p.add_run(label)
        run_l.bold = True
        run_l.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        run_v = p.add_run(val)
        run_v.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph() # Spacer

    current_list_paragraph = None

    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue

        # Skip main title block since we added it manually at the top
        if line_str.startswith("# ") and "Internship Daily Task Diary" in line_str:
            continue
        if line_str.startswith("**Duration**") or line_str.startswith("**Role**") or line_str.startswith("**Project**"):
            continue
        if line_str == "---":
            continue

        # Heading 2 (Weeks)
        if line_str.startswith("## "):
            week_title = line_str.replace("## ", "").replace("📅 ", "").strip()
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(week_title)
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a) # Navy blue
            current_list_paragraph = None
            continue

        # Bullet List Items (Days)
        if line_str.startswith("* ") or line_str.startswith("- "):
            # Clean list tags and bold wrappers
            item_text = re.sub(r'^[\*\-\s]+', '', line_str)
            
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            # Parse Day X block: "**Day X (Date)**: Description text"
            match = re.match(r'^\*\*(.*?)\*\*:(.*)', item_text)
            if match:
                day_prefix = match.group(1).strip()
                desc_text = match.group(2)
                
                # Add Day block (Bold)
                run_day = p.add_run(day_prefix + ":")
                run_day.bold = True
                run_day.font.name = 'Arial'
                run_day.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
                
                # Parse description for embedded file paths and code wrappers
                # Replace backticks or link formats
                desc_clean = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', desc_text)
                desc_clean = desc_clean.replace("`", "")
                
                # Add Description (Regular)
                run_desc = p.add_run(desc_clean)
                run_desc.font.name = 'Arial'
                run_desc.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            else:
                # Fallback if parsing fails
                item_clean = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', item_text).replace("`", "")
                run_item = p.add_run(item_clean)
                run_item.font.name = 'Arial'
                run_item.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            
            continue

        # Default fallback paragraph
        p = doc.add_paragraph()
        run = p.add_run(line_str)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.save(docx_path)
    print("Word document generated successfully!")

if __name__ == "__main__":
    create_diary_docx()
